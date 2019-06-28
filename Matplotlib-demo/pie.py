# -*- coding: utf-8 -*-
import matplotlib.pyplot as plt
from io import BytesIO
from matplotlib import font_manager
# 字体大小
plt.rcParams['font.size'] = 30
my_font = font_manager.FontProperties(fname="/home/inhoo/gitdir/IO-USE-Python/Matplotlib-demo/STKAITI.TTF")
# font = {'family' : 'SimHei',
#         'weight' : 'bold',
#         'size'   : '16'}
# plt.rc('font', **font)  # pass in the font dict as kwargs
# plt.rc('axes',unicode_minus=False)
# Pie chart, where the slices will be ordered and plotted counter-clockwise:
labels = u'阿萨德', u'阿萨德', u'阿萨德', u'阿萨德'
sizes = [0, 1, 1, 1]
explode = (0, 0, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')
colors = ['red','red','red','red','red','red','red']
fig1, ax1 = plt.subplots()
pie = ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.2f%%',colors=colors,
        shadow=False, startangle=90)
for font in pie[1]:
    font.set_fontproperties(my_font)
ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
# plt.savefig('test2.jpg')
buf = BytesIO()
fig1.savefig(buf, format="png",dpi=500)
# plt.show()
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')

document.add_picture(buf, width=Inches(2.25))

document.add_page_break()

document.save('demo1.docx')
