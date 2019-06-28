# -*- coding: utf-8 -*-


from docx import Document
from docx.section import Section
from docx.shared import Inches
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET


document = Document()
#
# for i in document.paragraphs:
#     print i.text
#     if '{{dep}}' in i.text:
#
#         i.text = i.text.replace('{{dep}}', u'')
#         run = i.add_run(u'中国移动')
#         run.font.size = Pt(26)
#         run.bold = True
#         run.font.name = u'微软雅黑'
#         run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
table.add_column(200)
table.add_row()
table.add_row()
table.rows[0].cells[0].text = '1'
# row.cells[0].merge(row.cells[-1])
document.save('demo.docx')